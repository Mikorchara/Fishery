"""Unify thesis title: body text uses 视频采集与传输 -> 视频分析系统 (matching cover)."""
import docx

SRC = "temp_thesis.docx"

doc = docx.Document(SRC)

REPLACEMENTS = [
    ("智慧渔业视频采集与传输系统原型", "智慧渔业视频分析系统原型"),
    ("smart fishery video acquisition and transmission system", "smart fishery video analysis system"),
]

changes = []

for p in doc.paragraphs:
    for run in p.runs:
        for old, new in REPLACEMENTS:
            if old in run.text:
                run.text = run.text.replace(old, new)
                changes.append(f"P: {old[:60]} -> {new[:60]}")

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    for old, new in REPLACEMENTS:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                            changes.append(f"T[{ti}]R{ri}C{ci}: {old[:60]} -> {new[:60]}")

doc.save(SRC)
print(f"Done. {len(changes)} changes:")
for c in changes:
    print(f"  {c}")
