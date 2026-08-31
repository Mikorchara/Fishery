"""Extract text and tables from docx files to UTF-8 text files."""
import sys
import docx

def extract_docx(path, out_path):
    with open(out_path, "w", encoding="utf-8") as out:
        doc = docx.Document(path)
        out.write("=== PARAGRAPHS ===\n")
        for i, p in enumerate(doc.paragraphs):
            style = p.style.name if p.style else ""
            text = p.text
            if text.strip():
                out.write(f"[{i}|{style}] {text}\n")
        out.write("\n=== TABLES ===\n")
        for ti, table in enumerate(doc.tables):
            out.write(f"--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---\n")
            for ri, row in enumerate(table.rows):
                cells = [cell.text[:200] for cell in row.cells]
                out.write(f"  Row {ri}: " + " | ".join(cells) + "\n")
        out.write("\n")

if __name__ == "__main__":
    for path in sys.argv[1:]:
        out_path = path.replace(".docx", ".txt")
        print(f"Extracting to {out_path}")
        extract_docx(path, out_path)
        print("Done.")
